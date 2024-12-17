from pypdf import PdfReader
from docx import Document
from .semantice_chunker import SemanticChunker
from llama_index.core.schema import Document as LlamaIndexDocument
from llama_index.core.schema import TextNode
from typing import List
from .nomic_embed import NomicEmbed
from numpy import ndarray
import json
from datetime import datetime
from utils import binary_array_to_base64, base64_to_binary_array, binary_quantized, float32_vector_to_base64, base64_to_float32_vector

def datetime_to_str(dt: datetime) -> str:
    if dt is None:
        return None
    return dt.strftime('%Y-%m-%d %H:%M:%S')

class D2D():
    """Class for converting documents to json form and then save it to disk in base64 format. And load back the base64 as json.

    Methods:
        _convert_pdf_to_json: Converts pdf file to json form.
        _convert_docx_to_json: Converts docx file to json form.
        convert_to_json: Converts file to json form.
        save_to_disk: Saves json data to disk.
        load_from_disk: Loads json data from disk.
    """
    
    def __init__(self):
        self.chunker = SemanticChunker()
        self.text_embedder = NomicEmbed()

    def __chunk_nodes_from_document(self, document: LlamaIndexDocument) -> List[TextNode]:
        """Chunks a document into a list of TextNodes.
        """
        nodes = self.chunker.chunk_nodes_from_documents([document])
        return nodes


    def _convert_pdf_to_dict(self, file_path: str) -> dict:
        """Converts pdf file to dict.
        """
        pdf = PdfReader(file_path)
        metadata = pdf.metadata
        metadata = {
            'title': metadata.title,
            'path': file_path,
            'author': metadata.author,
            'creator': metadata.creator,
            'creation_date': datetime_to_str(metadata.creation_date),
            'modification_date': datetime_to_str(metadata.modification_date),
            'producer': metadata.producer,
            'subject': metadata.subject,
        }
        text = ''
        for page in pdf.pages:
            text += page.extract_text()
        text = self.__chunk_nodes_from_document(LlamaIndexDocument(text=text))
        text_embed: list[ndarray] = []
        for chunk in text:
            quantized = binary_quantized(self.text_embedder._get_text_embedding(chunk.text))
            text_embed.append(quantized)
        
        id = 'user_' + str(hash(file_path))
        pdf_json = {
            'id': id,
            'metadata': metadata,
            'content': {
                f'{id}_text_{i}': {
                    'content': text[i].text,
                    'embedding': text_embed[i]
                    }
                for i in range(len(text))
            },
            'images': {
            }
        }
        return pdf_json
    
    def _convert_docx_to_dict(self, file_path: str) -> dict:
        """Converts docx file to json form.
        """
        doc = Document(file_path)
        metadata = doc.core_properties
        metadata = {
            'title': metadata.title,
            'path': file_path,
            'author': metadata.author,
            'category': metadata.category,
            'subject': metadata.subject,
            'comments': metadata.comments,
            'content_status': metadata.content_status,
            'created': datetime_to_str(metadata.created),
            'modified': datetime_to_str(metadata.modified),
            'last_modified_by': metadata.last_modified_by,
            'revision': metadata.revision,
        }
        text = ''
        for para in doc.paragraphs:
            text += para.text + '\n'

        text = self.__chunk_nodes_from_document(LlamaIndexDocument(text=text))
        text_embed: list[ndarray] = []
        for chunk in text:
            quantized = binary_quantized(self.text_embedder._get_text_embedding(chunk.text))
            text_embed.append(quantized)

        id = 'user_' + str(hash(file_path))
        doc_json = {
            'id': id,
            'metadata': metadata,
            'content': {
                f'{id}_text_{i}': {
                    'content': text[i].text,
                    'embedding': text_embed[i]
                    }
                for i in range(len(text))
            },
            'images': {
            }
        }
        return doc_json
    
    def convert_to_dict(self, file_path: str) -> dict | None:
        """Converts file to json form.
        """
        if file_path.endswith('.pdf'):
            return self._convert_pdf_to_dict(file_path)
        elif file_path.endswith('.docx'):
            return self._convert_docx_to_dict(file_path)
        else:
            print('Error: Unsupported file format.')
            return None
        
    def save_to_disk(self, file_path: str, data: dict) -> None:
        """Saves json data to disk.
        """
        for id in data["content"]:
            emb = data["content"][id]["embedding"]
            base64 = binary_array_to_base64(emb)
            data["content"][id]["embedding"] = base64
        for id in data["images"]:
            emb = data["images"][id]["embedding"]
            base64 = float32_vector_to_base64(emb)
            data["images"][id]["embedding"] = base64
        with open(file_path, 'w') as f:
            f.write(json.dumps(data, indent=4))

    def load_from_disk(self, file_path:str) -> dict:
        """Loads json data from disk.
        """
        with open(file_path, 'r') as f:
            data = json.load(f)
        for id in data["content"]:
            base64 = data["content"][id]["embedding"]
            emb = base64_to_binary_array(base64)
            data["content"][id]["embedding"] = emb
        for id in data["images"]:
            base64 = data["images"][id]["embedding"]
            emb = base64_to_float32_vector(base64)
            data["images"][id]["embedding"] = emb
        return data
